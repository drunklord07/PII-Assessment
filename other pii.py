import re
from docx import Document
from docx.shared import RGBColor

# Define the input file
input_file = "input.txt"

# Updated regex patterns for PII
pii_patterns = {
    "PAN": r"(?<![A-Z0-9])[A-Z]{5}[0-9]{4}[A-Z](?![A-Z0-9])",
    "Email": r"(?<![\w])[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,10}(?![\w])",
    "Mobile": r"(?<!\d)(?:\+91[\-\s]?|91[\-\s]?|91|0)?[6-9]\d{9}(?!\d)",
    "UPI": r"(?<![\w])[a-zA-Z0-9.\-_]{2,256}@[a-zA-Z0-9]{2,64}(?![\w])"   # UPI without dot after @
}

# Address keywords
address_keys = [
    "address", "full_address", "complete_address", "residential_address", "permanent_address",
    "current_address", "correspondence_address", "present_address", "mailing_address",
    "billing_address", "shipping_address", "registered_address", "home_address",
    "office_address", "work_address", "business_address", "shop_address", "delivery_address", "native_address",
    "house_no", "building_name", "flat_no", "apartment", "door_number", "plot_no", "block",
    "floor", "tower", "unit_number", "address_line1", "address_line2", "street", "street_name",
    "road", "lane", "area", "locality", "colony", "sector", "village", "district",
    "taluk", "mandal", "tehsil", "municipality", "town", "city", "state", "region",
    "zone", "division", "province", "pincode", "pin", "postal_code", "zip", "zip_code",
    "location", "geo_location", "place", "addr", "addr1", "addr2"
]

# Create a Word document for each PII type
documents = {pii: Document() for pii in pii_patterns.keys()}
documents["Address"] = Document()

# Match counters
match_counts = {pii: 0 for pii in pii_patterns.keys()}
match_counts["Address"] = 0

# Line counter
total_lines_scanned = 0

# Process input
with open(input_file, "r", encoding="utf-8") as file:
    for line_number, line in enumerate(file, start=1):
        total_lines_scanned += 1

        # Regex-based PII detection
        for pii_type, pattern in pii_patterns.items():
            for match in re.finditer(pattern, line):
                matched_text = match.group()
                para = documents[pii_type].add_paragraph(f"Line {line_number}: ")

                # Split text for highlighting
                start, end = match.span()
                before = line[:start]
                after = line[end:]

                para.add_run(before)

                match_run = para.add_run(matched_text)
                match_run.font.color.rgb = RGBColor(255, 0, 0)

                para.add_run(after.strip())
                match_counts[pii_type] += 1

        # Address keyword detection
        lowered = line.lower()
        for keyword in address_keys:
            if keyword.lower() in lowered:
                para = documents["Address"].add_paragraph(f"Line {line_number}: ")

                highlighted = re.sub(f"(?i)({re.escape(keyword)})", r"<<\1>>", line)
                parts = highlighted.split("<<")
                for part in parts:
                    if ">>" in part:
                        keyword_text, rest = part.split(">>", 1)
                        run = para.add_run(keyword_text)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        para.add_run(rest)
                    else:
                        para.add_run(part)
                match_counts["Address"] += 1
                break  # Avoid duplicate matches in same line

# Save all Word documents
for pii_type, document in documents.items():
    filename = f"{pii_type}_matches.docx"
    document.save(filename)

# Print summary
print("\n🔍 PII Scan Summary:")
print(f"- Total lines scanned: {total_lines_scanned}")
for pii_type, count in match_counts.items():
    print(f"- {pii_type}: {count} matches found")
print(f"\n✅ Done! Word files saved for each PII type.")
