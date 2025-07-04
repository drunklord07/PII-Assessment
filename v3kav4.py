import os
import re
from multiprocessing import Pool, cpu_count
from docx import Document
from docx.shared import RGBColor
from tqdm import tqdm

# --- CONFIGURATION ---
INPUT_FILE = "input.txt"
OUTPUT_DIR = "output"
CHUNK_SIZE = 1000
SAVE_BATCH_SIZE = 100

# --- PII REGEX PATTERNS ---
pii_patterns = {
    "PAN": r"(?<![A-Z0-9])[A-Z]{5}[0-9]{4}[A-Z](?![A-Z0-9])",
    "Email": r"(?<![\w])[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,10}(?![\w])",
    "Mobile": r"(?<!\d)(?:\+91[\-\s]?|91[\-\s]?|91|0)?[6-9]\d{9}(?!\d)",
    "UPI": r"(?<![\w])[a-zA-Z0-9.\-_]{2,256}@[a-zA-Z]{2,64}(?![\w])",
    "MAC": r"(?:[0-9A-Fa-f]{2}[:-]){5}(?:[0-9A-Fa-f]{2})",
    "IP": r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b",
    "Coordinates": r"-?\d{1,3}\.\d+,\s*-?\d{1,3}\.\d+",
    "CardNumber": r"(?:\d{4}[-\s]?){3}\d{4}|\d{15,16}",
    "GSTIN": r"\d{2}[A-Z]{5}\d{4}[A-Z]{1}[A-Z\d]{1}[Z]{1}[A-Z\d]{1}",
    "DLNumber": r"[A-Z]{2}\d{2}[-\s]?\d{4}\d{7}",
    "VoterID": r"[A-Z]{3}[0-9]{7}"
}
compiled_pii = {k: re.compile(v) for k, v in pii_patterns.items()}

# --- KEYWORD GROUPS ---
keyword_groups = {
    "Address": ["address", "full address", "complete address", "residential address", "permanent address",
                "locality", "pincode", "postal code", "zip", "zip code", "city", "state"],
    "Name": ["name"],
    "DOB": ["date of birth", "dob", "birthdate", "born on"],
    "AccountNumber": ["account number", "acc number", "bank account", "account no", "a/c no"],
    "CustomerID": ["customer id", "cust id", "customer number"],
    "SensitiveHints": ["national id", "identity card", "proof of identity", "document number"],
    "InsurancePolicy": ["insurance number", "policy number", "insurance id"]
}

def keyword_to_pattern(keyword):
    return re.escape(keyword).replace(r'\ ', r'[\s._-]*')

compiled_keywords = {
    cat: [re.compile(keyword_to_pattern(k), re.IGNORECASE) for k in keys]
    for cat, keys in keyword_groups.items()
}

# --- SCANNING FUNCTION ---
def process_chunk(args):
    start_line, lines = args
    results = {k: [] for k in list(compiled_pii) + list(compiled_keywords)}
    for idx, line in enumerate(lines):
        line_num = start_line + idx + 1
        lowered = line.lower()

        for pii_type, pattern in compiled_pii.items():
            for match in pattern.finditer(line):
                value = match.group()
                if pii_type == "IP":
                    octets = value.split('.')
                    if (octets[0] == '10' or
                        (octets[0] == '172' and 16 <= int(octets[1]) <= 31) or
                        (octets[0] == '192' and octets[1] == '168') or
                        value == '127.0.0.1'):
                        continue
                results[pii_type].append((line_num, line.strip(), match.span(), value))

        for cat, patterns in compiled_keywords.items():
            for pattern in patterns:
                match = pattern.search(lowered)
                if match:
                    results[cat].append((line_num, line.strip(), None, match.group()))
                    break
    return results

# --- MERGE RESULTS ---
def merge_results(partials):
    merged = {k: [] for k in list(compiled_pii) + list(compiled_keywords)}
    for part in partials:
        for k, v in part.items():
            merged[k].extend(v)
    return merged

# --- SAVE INDIVIDUAL DOCX FILE ---
def save_docx_batch(category, batch_index, items):
    folder = os.path.join(OUTPUT_DIR, category)
    os.makedirs(folder, exist_ok=True)
    filename = os.path.join(folder, f"{category}_{batch_index + 1}.docx")

    doc = Document()
    for line_num, text, span, match_text in items:
        para = doc.add_paragraph(f"Line {line_num}: ")
        if span:
            start, end = span
            para.add_run(text[:start])
            match = para.add_run(match_text)
            match.font.color.rgb = RGBColor(255, 0, 0)
            para.add_run(text[end:])
        else:
            match = para.add_run(text)
            match.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(filename)

# --- SAVE RESULTS SEQUENTIALLY ---
def save_results(results):
    print("\n📝 Saving Word files...")
    jobs = []
    for category, items in results.items():
        if not items:
            continue
        for i in range(0, len(items), SAVE_BATCH_SIZE):
            batch = items[i:i + SAVE_BATCH_SIZE]
            jobs.append((category, i // SAVE_BATCH_SIZE, batch))

    for args in tqdm(jobs, desc="📄 Saving DOCX Files", unit="file"):
        save_docx_batch(*args)

# --- MAIN EXECUTION ---
def main():
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    chunks = [(i, lines[i:i + CHUNK_SIZE]) for i in range(0, len(lines), CHUNK_SIZE)]

    print(f"🔍 Scanning {len(lines)} lines using {cpu_count()} cores...\n")
    with Pool(cpu_count()) as pool:
        partials = list(tqdm(
            pool.imap(process_chunk, chunks),
            total=len(chunks),
            desc="🔍 Scanning Chunks",
            unit="chunk"
        ))

    final_results = merge_results(partials)
    save_results(final_results)

    print("\n📊 PII Scan Summary:")
    for k, v in final_results.items():
        print(f"- {k}: {len(v)} matches")
    print(f"\n✅ Done! Files saved in '{OUTPUT_DIR}/<PII>/' folders.")

if __name__ == "__main__":
    main()
