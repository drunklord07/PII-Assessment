import re
import os
from docx import Document
from docx.shared import RGBColor

# Define the input file
input_file = "input.txt"

# Utility function to create regex patterns allowing _ . - spaces
def keyword_to_pattern(keyword):
    return re.sub(r"\s+", r"[\s._-]*", re.escape(keyword))

# Updated regex patterns for PII
pii_patterns = {
    "PAN": r"(?<![A-Z0-9])[A-Z]{5}[0-9]{4}[A-Z](?![A-Z0-9])",
    "Email": r"(?<![\w])[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,10}(?![\w])",
    "Mobile": r"(?<!\d)(?:\+91[\-\s]?|91[\-\s]?|91|0)?[6-9]\d{9}(?!\d)",
    "UPI": r"(?<![\w])[a-zA-Z0-9.\-_]{2,256}@[a-zA-Z0-9]{2,64}(?![\w])",
    "MAC": r"(?:[0-9A-Fa-f]{2}[:-]){5}(?:[0-9A-Fa-f]{2})",
    "IP": r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b",
    "Coordinates": r"-?\d{1,3}\.\d+,\s*-?\d{1,3}\.\d+",
    "CardNumber": r"(?:\d{4}[-\s]?){3}\d{4}|\d{15,16}",
    "GSTIN": r"\d{2}[A-Z]{5}\d{4}[A-Z]{1}[A-Z\d]{1}[Z]{1}[A-Z\d]{1}",
    "DLNumber": r"[A-Z]{2}\d{2}[-\s]?\d{4}\d{7}",
    "VoterID": r"[A-Z]{3}[0-9]{7}"
}

# Keyword categories
address_keys = [
    "address", "full address", "complete address", "residential address", "permanent address",
    "current address", "correspondence address", "present address", "mailing address",
    "billing address", "shipping address", "registered address", "home address",
    "office address", "work address", "business address", "shop address", "delivery address", "native address",
    "house no", "building name", "flat no", "apartment", "door number", "plot no", "block",
    "floor", "tower", "unit number", "address line1", "address line2", "street", "street name",
    "road", "lane", "area", "locality", "colony", "sector", "village", "district",
    "taluk", "mandal", "tehsil", "municipality", "town", "city", "state", "region",
    "zone", "division", "province", "pincode", "pin", "postal code", "zip", "zip code",
    "location", "geo location", "place", "addr", "addr1", "addr2"
]

name_keys = ["name"]

dob_keys = [
    "dob", "date of birth", "birth date", "d.o.b", "birthdate", "dateofbirth",
    "birth day", "birth", "born on"
]

account_keys = [
    "account number", "acc number", "account no", "account_no", "acc_no",
    "bank account", "bank account number", "acct number", "a/c number", "a/c no",
    "accountnum", "accountnumbr", "account", "account id",
    "beneficiary account", "beneficiary account number", "beneficiary acc", "beneficiary acct",
    "credited to account", "debited from account",
    "receiving account", "sender account", "payee account", "receiver account",
    "to account", "from account"
]

customer_keys = [
    "customer id", "cust id", "customerid", "custid",
    "customer number", "cust number", "customer no", "cust no",
    "customeridnumber", "custidnumber"
]

sensitive_keys = [
    "national id", "national identification number",
    "natl id", "natl_id",
    "document number", "doc number",
    "document id", "document_id",
    "doc id", "doc_id",
    "poi", "poa", "id proof", "identity document", "identity no",
    "identity card", "identification card",
    "proof of identity", "proof of address",
    "address proof"
]

insurance_keys = [
    "insurance", "insurance number", "insurance policy", "insurance id", "insuranceid", "insurance no",
    "policy number", "policy no", "policy id", "policyid", "ins id"
]

# Create Word documents
documents = {pii: Document() for pii in pii_patterns.keys()}
documents["Address"] = Document()
documents["Name"] = Document()
documents["DOB"] = Document()
documents["AccountNumber"] = Document()
documents["CustomerID"] = Document()
documents["SensitiveHints"] = Document()
documents["InsurancePolicy"] = Document()

# Match counters
match_counts = {pii: 0 for pii in pii_patterns.keys()}
match_counts.update({
    "Address": 0,
    "Name": 0,
    "DOB": 0,
    "AccountNumber": 0,
    "CustomerID": 0,
    "SensitiveHints": 0,
    "InsurancePolicy": 0
})

# Create output directory
output_dir = "output"
os.makedirs(output_dir, exist_ok=True)

# Process input
total_lines_scanned = 0
with open(input_file, "r", encoding="utf-8") as file:
    for line_number, line in enumerate(file, start=1):
        total_lines_scanned += 1
        lowered = line.lower()

        # Regex PII detection
        for pii_type, pattern in pii_patterns.items():
            for match in re.finditer(pattern, line):
                matched_text = match.group()

                # For IP addresses, skip private/internal
                if pii_type == "IP":
                    octets = matched_text.split('.')
                    if (octets[0] == '10' or
                        (octets[0] == '172' and 16 <= int(octets[1]) <= 31) or
                        (octets[0] == '192' and octets[1] == '168') or
                        matched_text == '127.0.0.1'):
                        continue

                para = documents[pii_type].add_paragraph(f"Line {line_number}: ")
                start, end = match.span()
                before = line[:start]
                after = line[end:]

                para.add_run(before)
                match_run = para.add_run(matched_text)
                match_run.font.color.rgb = RGBColor(255, 0, 0)
                para.add_run(after.strip())
                match_counts[pii_type] += 1

        # Keyword groups
        keyword_groups = [
            ("Address", address_keys),
            ("Name", name_keys),
            ("DOB", dob_keys),
            ("AccountNumber", account_keys),
            ("CustomerID", customer_keys),
            ("SensitiveHints", sensitive_keys),
            ("InsurancePolicy", insurance_keys)
        ]

        for pii_type, keys in keyword_groups:
            for keyword in keys:
                pattern = keyword_to_pattern(keyword)
                if re.search(pattern, lowered, re.IGNORECASE):
                    para = documents[pii_type].add_paragraph(f"Line {line_number}: ")

                    highlighted = re.sub(f"(?i)({pattern})", r"<<\1>>", lowered)
                    parts = highlighted.split("<<")
                    for part in parts:
                        if ">>" in part:
                            keyword_text, rest = part.split(">>", 1)
                            run = para.add_run(keyword_text)
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            para.add_run(rest)
                        else:
                            para.add_run(part)
                    match_counts[pii_type] += 1
                    break

# Save all Word documents
for pii_type, document in documents.items():
    filename = f"{pii_type}_matches.docx"
    document.save(os.path.join(output_dir, filename))

# Print summary
print("\n🔍 PII Scan Summary:")
print(f"- Total lines scanned: {total_lines_scanned}")
for pii_type, count in match_counts.items():
    print(f"- {pii_type}: {count} matches found")
print(f"\n✅ Done! Word files saved inside the '{output_dir}' folder.")
